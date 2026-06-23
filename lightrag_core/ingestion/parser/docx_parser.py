"""DOCX document parser implementation."""

import logging
from pathlib import Path

from lightrag_core.ingestion.parser import BaseParser

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Parser for Word documents (.docx).

    Uses python-docx to extract text from paragraphs and tables.
    """

    SUPPORTED_EXTENSIONS = {".docx"}

    def __init__(self) -> None:
        try:
            from docx import Document  # noqa: F401

            self._available = True
        except ImportError:
            logger.warning("python-docx not installed — Word parsing unavailable")
            self._available = False

    def parse(self, file_path: str | Path) -> str:
        if not self._available:
            raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        return "\n\n".join(parts)

    def supports(self, file_path: str | Path) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
