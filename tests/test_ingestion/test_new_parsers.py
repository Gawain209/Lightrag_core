"""Tests for new document parsers: DOCX, CSV, JSON, HTML."""

import json
import tempfile
from pathlib import Path

import pytest

from lightrag_core.ingestion.parser.docx_parser import WordParser
from lightrag_core.ingestion.parser.csv_parser import CSVParser
from lightrag_core.ingestion.parser.json_parser import JSONParser
from lightrag_core.ingestion.parser.html_parser import HTMLParser
from lightrag_core.ingestion.parser.doc_parser import DocParser
from lightrag_core.ingestion.parser.xlsx_parser import XlsxParser


class TestWordParser:
    def test_parse_docx(self) -> None:
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        try:
            doc = Document()
            doc.add_paragraph("Paragraph one.")
            doc.add_paragraph("Paragraph two.")
            doc.save(temp_path)

            parser = WordParser()
            content = parser.parse(temp_path)
            assert "Paragraph one." in content
            assert "Paragraph two." in content
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_supports(self) -> None:
        parser = WordParser()
        assert parser.supports("file.docx") is True
        assert parser.supports("file.doc") is False
        assert parser.supports("file.txt") is False


class TestCSVParser:
    def test_parse_csv(self) -> None:
        csv_content = "name,age,city\nAlice,30,NYC\nBob,25,LA"
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".csv", delete=False, encoding="utf-8"
        ) as f:
            f.write(csv_content)
            temp_path = f.name

        try:
            parser = CSVParser()
            content = parser.parse(temp_path)
            assert "name: Alice" in content
            assert "age: 25" in content
            assert "city: LA" in content
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_parse_empty_csv(self) -> None:
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".csv", delete=False, encoding="utf-8"
        ) as f:
            f.write("")
            temp_path = f.name

        try:
            parser = CSVParser()
            content = parser.parse(temp_path)
            assert content == ""
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_supports(self) -> None:
        parser = CSVParser()
        assert parser.supports("data.csv") is True
        assert parser.supports("data.json") is False


class TestJSONParser:
    def test_parse_json_object(self) -> None:
        data = {"name": "Alice", "age": 30, "city": "NYC"}
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".json", delete=False, encoding="utf-8"
        ) as f:
            json.dump(data, f)
            temp_path = f.name

        try:
            parser = JSONParser()
            content = parser.parse(temp_path)
            assert "name: Alice" in content
            assert "age: 30" in content
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_parse_json_array(self) -> None:
        data = [
            {"product": "A", "price": 10},
            {"product": "B", "price": 20},
        ]
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".json", delete=False, encoding="utf-8"
        ) as f:
            json.dump(data, f)
            temp_path = f.name

        try:
            parser = JSONParser()
            content = parser.parse(temp_path)
            assert "product: A" in content
            assert "price: 20" in content
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_parse_nested_json(self) -> None:
        data = {"user": {"name": "Alice", "address": {"city": "NYC"}}}
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".json", delete=False, encoding="utf-8"
        ) as f:
            json.dump(data, f)
            temp_path = f.name

        try:
            parser = JSONParser()
            content = parser.parse(temp_path)
            assert "user.address.city: NYC" in content
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_supports(self) -> None:
        parser = JSONParser()
        assert parser.supports("data.json") is True
        assert parser.supports("data.csv") is False


class TestHTMLParser:
    def test_parse_html(self) -> None:
        try:
            from bs4 import BeautifulSoup  # noqa: F401
        except ImportError:
            pytest.skip("beautifulsoup4 not installed")

        html = """<html><body>
            <h1>Title</h1>
            <p>Paragraph text here.</p>
            <script>console.log('hidden')</script>
            <style>body { color: red; }</style>
        </body></html>"""
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".html", delete=False, encoding="utf-8"
        ) as f:
            f.write(html)
            temp_path = f.name

        try:
            parser = HTMLParser()
            content = parser.parse(temp_path)
            assert "Title" in content
            assert "Paragraph text here." in content
            assert "console.log" not in content  # script stripped
            assert "color: red" not in content  # style stripped
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_supports(self) -> None:
        parser = HTMLParser()
        assert parser.supports("page.html") is True
        assert parser.supports("page.htm") is True
        assert parser.supports("page.txt") is False


class TestXlsxParser:
    def test_parse_xlsx(self) -> None:
        try:
            from openpyxl import Workbook  # noqa: F401
        except ImportError:
            pytest.skip("openpyxl not installed")

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            temp_path = f.name
        try:
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws.title = "Sheet1"
            ws.append(["name", "age", "city"])
            ws.append(["Alice", 30, "NYC"])
            ws.append(["Bob", 25, "LA"])
            wb.save(temp_path)

            parser = XlsxParser()
            content = parser.parse(temp_path)
            assert "name: Alice" in content
            assert "age: 25" in content
            assert "city: LA" in content
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_supports(self) -> None:
        parser = XlsxParser()
        assert parser.supports("data.xlsx") is True
        assert parser.supports("data.xls") is False
        assert parser.supports("data.csv") is False


class TestDocParser:
    def test_parse_invalid_doc(self) -> None:
        try:
            import olefile  # noqa: F401
        except ImportError:
            pytest.skip("olefile not installed")

        parser = DocParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/path.doc")

        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            f.write(b"not a valid OLE2 compound document")
            temp_path = f.name
        try:
            with pytest.raises(RuntimeError):
                parser.parse(temp_path)
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_supports(self) -> None:
        parser = DocParser()
        assert parser.supports("file.doc") is True
        assert parser.supports("file.docx") is False
        assert parser.supports("file.txt") is False
